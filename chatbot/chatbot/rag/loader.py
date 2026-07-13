import json
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from langchain_core.documents import Document
from pypdf import PdfReader

from rag.progress import log_progress, progress_bar


class DocumentLoader:
    """Loads PDF and JSON files from data directories."""

    def load_pdfs(self, pdf_dir: Path, show_progress: bool = False) -> list[Document]:
        documents: list[Document] = []
        pdf_paths = sorted(pdf_dir.glob("*.pdf"))
        path_iter = (
            progress_bar(pdf_paths, desc="Loading PDFs", unit="file")
            if show_progress
            else pdf_paths
        )
        for pdf_path in path_iter:
            reader = PdfReader(str(pdf_path))
            text_parts = []
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")

            full_text = "\n\n".join(text_parts)
            if full_text.strip():
                documents.append(
                    Document(
                        page_content=full_text,
                        metadata={
                            "source": pdf_path.name,
                            "doc_type": "pdf",
                            "path": str(pdf_path),
                        },
                    )
                )
        return documents

    def load_docx_files(self, doc_dir: Path, show_progress: bool = False) -> list[Document]:
        documents: list[Document] = []
        docx_paths = sorted(doc_dir.glob("*.docx"))
        path_iter = (
            progress_bar(docx_paths, desc="Loading DOCX", unit="file")
            if show_progress
            else docx_paths
        )
        for docx_path in path_iter:
            doc = DocxDocument(str(docx_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            full_text = "\n\n".join(paragraphs)
            if full_text.strip():
                documents.append(
                    Document(
                        page_content=full_text,
                        metadata={
                            "source": docx_path.name,
                            "doc_type": "docx",
                            "path": str(docx_path),
                        },
                    )
                )
        return documents

    def _json_to_readable_text(self, data: object, prefix: str = "") -> str:
        """Recursively convert JSON into human-readable text for embedding."""
        if isinstance(data, dict):
            lines = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._json_to_readable_text(value, prefix + "  "))
                else:
                    lines.append(f"{prefix}{key}: {value}")
            return "\n".join(lines)
        if isinstance(data, list):
            lines = []
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{prefix}Item {i + 1}:")
                    lines.append(self._json_to_readable_text(item, prefix + "  "))
                else:
                    lines.append(f"{prefix}- {item}")
            return "\n".join(lines)
        return f"{prefix}{data}"

    def _json_records_to_documents(
        self, data: Any, json_path: Path
    ) -> list[Document]:
        """Create one Document per CPT-like record when the JSON is a list of dicts."""
        if not isinstance(data, list) or not all(isinstance(item, dict) for item in data):
            readable_text = self._json_to_readable_text(data)
            return [
                Document(
                    page_content=readable_text,
                    metadata={
                        "source": json_path.name,
                        "doc_type": "json",
                        "path": str(json_path),
                    },
                )
            ]

        documents: list[Document] = []
        for index, item in enumerate(data, start=1):
            item = dict(item)
            cpt_code = item.get("cpt_code")
            readable_text = self._json_to_readable_text(item)
            documents.append(
                Document(
                    page_content=readable_text,
                    metadata={
                        "source": json_path.name,
                        "doc_type": "json",
                        "path": str(json_path),
                        "record_index": index,
                        "cpt_code": str(cpt_code) if cpt_code is not None else "",
                    },
                )
            )
        return documents

    def load_json_files(self, json_dir: Path, show_progress: bool = False) -> list[Document]:
        documents: list[Document] = []
        json_paths = sorted(json_dir.glob("*.json"))
        if show_progress:
            log_progress(f"Found {len(json_paths)} JSON file(s) in {json_dir}")

        path_iter = (
            progress_bar(json_paths, desc="Loading JSON", unit="file")
            if show_progress
            else json_paths
        )

        for json_path in path_iter:
            with open(json_path, encoding="utf-8-sig") as f:
                data = json.load(f)
            file_docs = self._json_records_to_documents(data, json_path)
            documents.extend(file_docs)
            if not show_progress:
                print(f"  Loaded {json_path.name}: {len(file_docs)} document(s)")

        if show_progress:
            log_progress(
                f"JSON total: {len(documents)} document(s) from {len(json_paths)} file(s)"
            )
        else:
            print(f"JSON total: {len(documents)} document(s) from {len(json_paths)} file(s)")
        return documents

    def load_all(
        self, json_dir: Path, pdf_dir: Path, show_progress: bool = False
    ) -> list[Document]:
        docs: list[Document] = []
        docs.extend(self.load_json_files(json_dir, show_progress=show_progress))
        docs.extend(self.load_pdfs(pdf_dir, show_progress=show_progress))
        docs.extend(self.load_docx_files(pdf_dir, show_progress=show_progress))
        return docs